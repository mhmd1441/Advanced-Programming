from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUTPUT = "TaskManagementSystem_Documentation.docx"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False, color=None):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)
    for paragraph in cell.paragraphs:
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.line_spacing = 1.05
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def set_table_borders(table, color="D9E2EC"):
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = "w:{}".format(edge)
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "6")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def style_table(table, header_fill="1F4E79"):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    set_table_borders(table)
    for row_index, row in enumerate(table.rows):
        for cell in row.cells:
            cell.margin_top = Inches(0.06)
            cell.margin_bottom = Inches(0.06)
            cell.margin_left = Inches(0.08)
            cell.margin_right = Inches(0.08)
            if row_index == 0:
                set_cell_shading(cell, header_fill)
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True
                        run.font.color.rgb = RGBColor(255, 255, 255)
            elif row_index % 2 == 0:
                set_cell_shading(cell, "F8FAFC")


def add_heading(doc, text, level=1):
    paragraph = doc.add_heading(text, level=level)
    for run in paragraph.runs:
        run.font.color.rgb = RGBColor(31, 78, 121)
    return paragraph


def add_body_paragraph(doc, text):
    paragraph = doc.add_paragraph(text)
    paragraph.paragraph_format.space_after = Pt(7)
    paragraph.paragraph_format.line_spacing = 1.08
    return paragraph


def add_bullets(doc, items):
    for item in items:
        paragraph = doc.add_paragraph(style="List Bullet")
        paragraph.add_run(item)
        paragraph.paragraph_format.space_after = Pt(3)


def add_numbered(doc, items):
    for item in items:
        paragraph = doc.add_paragraph(style="List Number")
        paragraph.add_run(item)
        paragraph.paragraph_format.space_after = Pt(3)


def add_table(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    hdr_cells = table.rows[0].cells
    for index, header in enumerate(headers):
        set_cell_text(hdr_cells[index], header, bold=True, color=(255, 255, 255))
    for row in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            set_cell_text(cells[index], str(value))
    if widths:
        for row in table.rows:
            for index, width in enumerate(widths):
                row.cells[index].width = width
    style_table(table)
    doc.add_paragraph()
    return table


def setup_document():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)

    styles = doc.styles
    styles["Normal"].font.name = "Aptos"
    styles["Normal"].font.size = Pt(10.5)
    styles["Title"].font.name = "Aptos Display"
    styles["Title"].font.size = Pt(26)
    styles["Heading 1"].font.name = "Aptos Display"
    styles["Heading 1"].font.size = Pt(16)
    styles["Heading 2"].font.name = "Aptos Display"
    styles["Heading 2"].font.size = Pt(13)
    styles["Heading 3"].font.name = "Aptos"
    styles["Heading 3"].font.size = Pt(11.5)
    return doc


def build_doc():
    doc = setup_document()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(72)
    title_run = title.add_run("Task Management System")
    title_run.bold = True
    title_run.font.size = Pt(30)
    title_run.font.color.rgb = RGBColor(31, 78, 121)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.add_run("Project Documentation")
    subtitle_run.font.size = Pt(17)
    subtitle_run.font.color.rgb = RGBColor(80, 80, 80)

    meta = doc.add_table(rows=4, cols=2)
    meta.alignment = WD_TABLE_ALIGNMENT.CENTER
    meta_data = [
        ("Course", "Advanced Programming"),
        ("Project Type", ".NET Task Management System"),
        ("Backend", "ASP.NET Core Web API, EF Core Code First, PostgreSQL, ASP.NET Identity, JWT"),
        ("Frontend", "ASP.NET Core MVC consuming the API through HttpClient"),
    ]
    for row, (label, value) in zip(meta.rows, meta_data):
        set_cell_text(row.cells[0], label, bold=True)
        set_cell_text(row.cells[1], value)
    style_table(meta, header_fill="D9EAF7")
    for cell in meta.rows[0].cells:
        for run in cell.paragraphs[0].runs:
            run.font.color.rgb = RGBColor(0, 0, 0)

    doc.add_section(WD_SECTION_START.NEW_PAGE)

    add_heading(doc, "1. Project Overview", 1)
    add_body_paragraph(
        doc,
        "The Task Management System is a complete .NET-based application designed to simulate a company task workflow. "
        "It allows users to register, log in, create and view tasks, organize work by category, assign tasks to users, "
        "and control access based on company roles. The system contains a secure REST API backend and an ASP.NET Core MVC frontend."
    )
    add_body_paragraph(
        doc,
        "The project demonstrates RESTful API development, Entity Framework Core Code First database design, dependency injection, "
        "JWT authentication, ASP.NET Identity role management, and frontend-to-backend integration."
    )

    add_heading(doc, "2. Requirement Coverage", 1)
    add_table(
        doc,
        ["Requirement", "Implementation"],
        [
            ("Use .NET Core for backend development", "Implemented using ASP.NET Core Web API targeting .NET 8."),
            ("RESTful API endpoints for task management", "Implemented endpoints for tasks, categories, authentication, users, and role updates."),
            ("Entity Framework Core Code First", "ApplicationDbContext and migrations create Identity, TaskItems, and Categories tables."),
            ("Dependency Injection / IoC", "Controllers depend on service interfaces such as ITaskService and ICategoryService."),
            ("Authentication and authorization", "JWT login with ASP.NET Identity roles: Employee, Manager, and Admin."),
            ("Frontend or API testing tool", "MVC frontend is implemented and Swagger is configured for API testing."),
            ("Documentation", "This document describes setup, endpoints, roles, and database schema."),
        ],
        widths=[Inches(2.4), Inches(4.8)],
    )

    add_heading(doc, "3. Technology Stack", 1)
    add_bullets(
        doc,
        [
            "ASP.NET Core Web API for backend REST endpoints.",
            "ASP.NET Core MVC for the web user interface.",
            "Entity Framework Core Code First for database schema and migrations.",
            "PostgreSQL as the relational database provider.",
            "ASP.NET Identity for users, password hashing, and role storage.",
            "JWT Bearer authentication for protected API access.",
            "Swagger / OpenAPI for API exploration and testing.",
            "IHttpClientFactory and session storage in MVC for API integration.",
        ],
    )

    add_heading(doc, "4. Project Setup", 1)
    add_heading(doc, "4.1 Prerequisites", 2)
    add_bullets(doc, [".NET 8 SDK", "PostgreSQL server", "Visual Studio or another .NET IDE", "Entity Framework Core tools"])

    add_heading(doc, "4.2 Configuration", 2)
    add_body_paragraph(
        doc,
        "The API database connection string is configured in TaskManagementAPI/appsettings.json. "
        "The current local connection uses PostgreSQL on localhost with database name TaskManagmentSystemAdvancedProgramming."
    )
    add_table(
        doc,
        ["Setting", "Value"],
        [
            ("API URL", "http://localhost:5199"),
            ("MVC URL", "http://localhost:5193"),
            ("Swagger URL", "http://localhost:5199/swagger"),
            ("Database", "TaskManagmentSystemAdvancedProgramming"),
            ("Provider", "Npgsql.EntityFrameworkCore.PostgreSQL"),
        ],
        widths=[Inches(2.0), Inches(5.2)],
    )

    add_heading(doc, "4.3 Running the Application", 2)
    add_numbered(
        doc,
        [
            "Start PostgreSQL and confirm the connection string is correct.",
            "Apply EF Core migrations using: dotnet ef database update --project TaskManagementAPI.",
            "Run both TaskManagementAPI and TaskManagementMvc from Visual Studio as startup projects, or run them from terminal.",
            "Open the MVC application at http://localhost:5193.",
            "Open Swagger at http://localhost:5199/swagger for API testing.",
        ],
    )

    add_heading(doc, "5. Authentication and Role Model", 1)
    add_body_paragraph(
        doc,
        "The system uses ASP.NET Identity for account management. Roles are stored in the standard Identity tables AspNetRoles "
        "and AspNetUserRoles. New accounts are always registered as Employee by default, and only an Admin can change a user's role."
    )
    add_table(
        doc,
        ["Role", "Permissions"],
        [
            ("Employee", "Can create tasks for themselves and view only tasks assigned to them."),
            ("Manager", "Can view all tasks and change task assignments. Cannot manage users or fully edit task details."),
            ("Admin", "Can manage users and roles, manage categories, and create, edit, assign, or delete all tasks."),
        ],
        widths=[Inches(1.4), Inches(5.8)],
    )

    add_heading(doc, "6. API Endpoints", 1)
    add_heading(doc, "6.1 Authentication Endpoints", 2)
    add_table(
        doc,
        ["Method", "Endpoint", "Description", "Authorization"],
        [
            ("POST", "/api/Auth/register", "Register a new user as Employee.", "Public"),
            ("POST", "/api/Auth/login", "Authenticate user and return JWT token.", "Public"),
        ],
        widths=[Inches(0.8), Inches(1.8), Inches(3.3), Inches(1.2)],
    )

    add_heading(doc, "6.2 Task Endpoints", 2)
    add_table(
        doc,
        ["Method", "Endpoint", "Description", "Authorization"],
        [
            ("GET", "/api/Tasks", "Get visible tasks. Employees see assigned tasks only; managers/admins see all.", "Authenticated"),
            ("GET", "/api/Tasks/{id}", "Get a visible task by ID.", "Authenticated"),
            ("POST", "/api/Tasks", "Create a task. Employees create for themselves; admins may assign.", "Admin, Employee"),
            ("PUT", "/api/Tasks/{id}", "Fully update task details.", "Admin"),
            ("PUT", "/api/Tasks/{id}/assignment", "Change assigned user.", "Admin, Manager"),
            ("DELETE", "/api/Tasks/{id}", "Delete a task.", "Admin"),
        ],
        widths=[Inches(0.7), Inches(1.8), Inches(3.6), Inches(1.1)],
    )

    add_heading(doc, "6.3 Category Endpoints", 2)
    add_table(
        doc,
        ["Method", "Endpoint", "Description", "Authorization"],
        [
            ("GET", "/api/Categories", "Get all task categories.", "Authenticated"),
            ("GET", "/api/Categories/{id}", "Get category by ID.", "Authenticated"),
            ("POST", "/api/Categories", "Create category.", "Admin"),
            ("PUT", "/api/Categories/{id}", "Update category.", "Admin"),
            ("DELETE", "/api/Categories/{id}", "Delete category.", "Admin"),
        ],
        widths=[Inches(0.7), Inches(1.9), Inches(3.5), Inches(1.1)],
    )

    add_heading(doc, "6.4 User and Role Endpoints", 2)
    add_table(
        doc,
        ["Method", "Endpoint", "Description", "Authorization"],
        [
            ("GET", "/api/Users", "List users for assignment and admin management.", "Admin, Manager"),
            ("PUT", "/api/Users/{id}/role", "Update a user's role to Employee, Manager, or Admin.", "Admin"),
        ],
        widths=[Inches(0.7), Inches(2.1), Inches(3.4), Inches(1.0)],
    )

    add_heading(doc, "7. Database Schema", 1)
    add_body_paragraph(
        doc,
        "The database is generated using Entity Framework Core Code First migrations. ApplicationDbContext inherits from "
        "IdentityDbContext<ApplicationUser>, which creates the standard ASP.NET Identity tables in addition to project-specific tables."
    )
    add_heading(doc, "7.1 Main Tables", 2)
    add_table(
        doc,
        ["Table", "Purpose"],
        [
            ("AspNetUsers", "Stores registered users, including FirstName, LastName, JobTitle, Email, PhoneNumber, and CreatedAt."),
            ("AspNetRoles", "Stores system roles such as Employee, Manager, and Admin."),
            ("AspNetUserRoles", "Many-to-many relationship table connecting users to roles."),
            ("Categories", "Stores task categories with Name and Description."),
            ("TaskItems", "Stores task data including Title, Description, Status, Priority, DueDate, CreatedAt, CategoryId, and AssignedToUserId."),
        ],
        widths=[Inches(1.8), Inches(5.4)],
    )

    add_heading(doc, "7.2 Relationships", 2)
    add_bullets(
        doc,
        [
            "One Category can contain many TaskItems through TaskItem.CategoryId.",
            "One ApplicationUser can have many assigned TaskItems through TaskItem.AssignedToUserId.",
            "Users and roles are connected through the Identity AspNetUserRoles table.",
            "Identity also manages claims, logins, tokens, password hashes, and security metadata.",
        ],
    )

    add_heading(doc, "8. MVC Frontend", 1)
    add_body_paragraph(
        doc,
        "The MVC project provides the user-facing interface and communicates with the API using IHttpClientFactory. "
        "After login, the JWT token is stored in session and attached to protected API requests."
    )
    add_table(
        doc,
        ["Screen", "Purpose"],
        [
            ("Login", "Authenticates the user and stores the JWT token in session."),
            ("Register", "Creates a new Employee account."),
            ("Dashboard", "Entry point for tasks, categories, and admin user management."),
            ("Tasks", "Lists tasks according to role permissions and supports create, detail, edit, assign, and delete actions."),
            ("Categories", "Displays categories and allows admins to create, edit, or delete them."),
            ("Users", "Admin user-management page for changing roles."),
        ],
        widths=[Inches(1.6), Inches(5.6)],
    )

    add_heading(doc, "9. Security Notes", 1)
    add_bullets(
        doc,
        [
            "Users cannot choose their own role during registration.",
            "Role changes are restricted to Admin users.",
            "JWT tokens include user identity and role claims.",
            "Protected API endpoints require Authorization headers with Bearer tokens.",
            "Task visibility is filtered so Employees only see tasks assigned to them.",
        ],
    )

    add_heading(doc, "10. Conclusion", 1)
    add_body_paragraph(
        doc,
        "The Task Management System satisfies the final project requirements by combining a RESTful API, Code First database design, "
        "dependency injection, authentication, authorization, MVC frontend integration, Swagger testing, and clear project documentation. "
        "The role model reflects a realistic company workflow and demonstrates secure access control across backend and frontend layers."
    )

    doc.save(OUTPUT)


if __name__ == "__main__":
    build_doc()
